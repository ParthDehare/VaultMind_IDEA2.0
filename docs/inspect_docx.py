from docx import Document

def inspect_doc(filename):
    print(f"--- {filename} ---")
    doc = Document(filename)
    
    # Check paragraphs
    for i, p in enumerate(doc.paragraphs):
        if any(keyword in p.text for keyword in ["1,000,000", "1M", "Phase 1", "Phase 3", "Section A2", "Schema", "unified_score", "agent"]):
            print(f"P[{i}]: {p.text}".encode('ascii', 'replace').decode('ascii'))
            
    # Check tables
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            row_data = [cell.text for cell in row.cells]
            if any("score" in c.lower() or "1,000,000" in c or "schema" in c.lower() or "text" in c.lower() for c in row_data):
                print(f"Table[{t_idx}] Row[{r_idx}]: {row_data}".encode('ascii', 'replace').decode('ascii'))

inspect_doc('02_Data_Creation_Parameters.docx')
inspect_doc('VaultMind_Complete_Build_Guide.docx')
inspect_doc('05_Testing_Checklist.docx')
