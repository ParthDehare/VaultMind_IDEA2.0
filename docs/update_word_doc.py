import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(16)
    else:
        run.font.size = Pt(14)
    return p

def add_paragraph(doc, text, bold=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph()
    run = p.add_run('• ' + text)
    return p

doc = Document('VaultMind_Complete_Build_Guide.docx')

add_heading(doc, 'Phase 8 — Profile Detail View: Context-Sensitive Investigator UI', level=1)
add_paragraph(doc, "When an investigator clicks any employee card on the Employee Watch screen, the system renders one of two entirely different panel layouts depending on the employee's Unified Threat Score.")
add_paragraph(doc, "The goal of each view is different:", bold=True)
add_bullet(doc, "Forensic View (Score 90–100): Prove fraud and enable legal action.")
add_bullet(doc, "Baseline View (Score 0–40): Prove innocence and provide reassurance.")

add_heading(doc, 'HIGH-RISK PROFILE (Score: 90–100) — "The Forensic View"', level=2)
add_paragraph(doc, "Investigator Intent: Investigate & Take Action", bold=True)
add_paragraph(doc, "This view is shown when the Unified Threat Score is between 90 and 100. Every panel is purpose-built to help the FCU investigator build a legally defensible case. The layout uses dark-accented styling with red warning highlights throughout.")

add_heading(doc, 'Panel 1 — Threat Spike Graph (Time-Series)', level=3)
add_bullet(doc, "Library: Recharts <AreaChart> or <LineChart>")
add_bullet(doc, "X-axis: Last 30 days (daily tick marks)")
add_bullet(doc, "Y-axis: Unified Threat Score (0–100)")
add_bullet(doc, "Line rendered in a gradient from amber (#F59E0B) at baseline to crimson (#DC2626) at peak")
add_bullet(doc, "A vertical dashed annotation line marks the Anomaly Onset Date — the first day the score crossed the alert threshold")
add_bullet(doc, "Two behavioural patterns are distinguishable: Sudden Spike (one-day jump > 40 points) and Slow Boil (linear rise over 30+ days)")
add_bullet(doc, "Tooltip on hover shows: Date | Score | Dominant Agent")
add_paragraph(doc, "Purpose: Proves WHEN the anomaly began — critical for linking it to a real-world event.", color=RGBColor(29, 78, 216))

add_heading(doc, 'Panel 2 — FundFlow Network Graph (GNN View)', level=3)
add_bullet(doc, "Library: Cytoscape.js with cose-bilkent or dagre layout")
add_paragraph(doc, "RENDER CONDITION: This panel only renders if Agent 2 (FundFlow) or Agent 5 (NetworkIntel) contributed a score > 0. If neither fired, this panel is replaced with: 'No network anomaly detected by Agent 2.'", bold=True)
add_bullet(doc, "Node types: Employee (blue circle), Shell Company (red hexagon), Relative/Linked Account (orange diamond), External Bank (grey square)")
add_bullet(doc, "Edge colours: Green = normal flow | Red = flagged circular routing | Dashed = indirect/inferred link")
add_bullet(doc, "Fraudulent loop highlighted with pulsing red edge animation: Manager → Shell Company → Relative → Manager")
add_bullet(doc, "Node size encodes transaction volume (larger = more money moved)")
add_bullet(doc, "Clicking a node opens a tooltip: Entity Name | Role | Total Flow (₹) | First Seen")
add_paragraph(doc, "Purpose: Visually proves the Circular Routing pattern — the backbone of layering fraud.", color=RGBColor(29, 78, 216))

add_heading(doc, 'Panel 3 — Counterfactual Sliders (XAI / SHAP Interface)', level=3)
add_bullet(doc, "Library: Radix UI <Slider> components — debounced API call to /api/xai/counterfactual/{employee_id}")
add_bullet(doc, "Live score readout: Unified Threat Score updates in real time (debounced 300ms) as sliders are dragged")
add_bullet(doc, "Critical Threshold Line: A horizontal line at score = 60 shows exactly how far parameters need to shift before the employee drops out of CRITICAL status")
add_paragraph(doc, "Purpose: Answers the key legal question 'Why did the AI flag this person?'", color=RGBColor(29, 78, 216))

add_heading(doc, 'The Architecture of the "What-If" Slider', level=4)
add_paragraph(doc, "1. The Read-Only Rule (Backend Level)", bold=True)
add_paragraph(doc, "When the dashboard is opened, the FastAPI backend serves Core Banking logs in strictly Read-Only mode. VaultMind's backend has no APIs exposed to UPDATE or DELETE these logs. The original transaction records remain immutable in PostgreSQL.")

add_paragraph(doc, "2. The In-Memory Sandbox (React Level)", bold=True)
add_paragraph(doc, "As the FCU investigator adjusts the sliders, they are manipulating an in-memory sandbox in the browser's RAM, not the actual data.")
add_bullet(doc, "We maintain an actualData state (which remains fixed).")
add_bullet(doc, "We maintain a simulatedData state (which reacts to the slider).")
add_paragraph(doc, "When the investigator shifts the 'Time' slider from 2 AM to 2 PM, the React frontend sends a 'Dummy Payload' to the AI model querying: 'If the time was 2 PM, what would the score be?' The AI responds with the hypothetical score (e.g., 41), which the UI then displays. Crucially, the actual database retains the 2 AM timestamp and the original 96 threat score.")

add_paragraph(doc, "3. The 'Audit Trail' (Taking Action)", bold=True)
add_paragraph(doc, "If the investigator determines after using the sliders that the alert is genuinely a 'False Positive' (i.e., not fraud), they can leave the slider in place and click the [Mark as False Positive] button. This will log their decision and trigger the necessary HITL adjustments.")

add_heading(doc, 'Panel 4 — Action Block', level=3)
add_bullet(doc, "PRIMARY BUTTON: [Generate STR Evidence] (Red, triggers Agent 7 pipeline for SHA-256 hash, PDF, STR draft)")
add_bullet(doc, "SECONDARY BUTTON: [Mark as False Positive] (Outlined grey, opens confirmation dialog for justification)")

add_heading(doc, 'NORMAL PROFILE (Score: 0–40) — "The Baseline View"', level=2)
add_paragraph(doc, "Investigator Intent: Reassurance — Prove why this employee is safe", bold=True)
add_paragraph(doc, "This view is shown when the Unified Threat Score is between 0 and 40. The design language shifts from red/dark-threat to green/calm. No fraud-investigation panels are shown.")

add_heading(doc, 'Panel 1 — Peer Comparison Radar Chart', level=3)
add_bullet(doc, "Library: Recharts <RadarChart> with two overlapping datasets")
add_bullet(doc, "Six radar axes: Daily Transaction Volume, Working Hours, Record Access Rate, Geographic Consistency, After-Hours Activity, Peer Deviation Score")
add_bullet(doc, "Blue fill = This employee's normalised scores")
add_bullet(doc, "Grey fill = Median of all ~50 clerks in the same branch peer cluster")

add_heading(doc, 'Panel 2 — Activity Heatmap (GitHub-style Contribution Graph)', level=3)
add_bullet(doc, "Layout: 52 columns × 7 rows (1 year rolling view)")
add_bullet(doc, "Any cell outside the 6 AM–8 PM window is capped at minimum colour — visually proving the employee is inactive at night")

add_heading(doc, 'Panel 3 — Threat Trend (Flatline Graph)', level=3)
add_bullet(doc, "X-axis: Last 30 days, Y-axis: Capped at 50 to visually emphasise low range")
add_bullet(doc, "Line appears flat, oscillating between 15 and 25 with no meaningful trend")

add_heading(doc, 'Panel 4 — Network Status Message', level=3)
add_paragraph(doc, "Design: A single clean info card with a green left border. The Cytoscape.js GNN graph is NOT rendered.")
add_bullet(doc, "System Health: Normal")
add_bullet(doc, "No suspicious internal or external network connections detected.")

doc.save('VaultMind_Complete_Build_Guide.docx')
print("Successfully appended Profile Detail View to VaultMind_Complete_Build_Guide.docx")
