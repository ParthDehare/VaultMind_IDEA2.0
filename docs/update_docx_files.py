import docx

# 1. Update 02_Data_Creation_Parameters.docx
doc1 = docx.Document('02_Data_Creation_Parameters.docx')

# Remove scores and add NLP columns in schema tables
for table in doc1.tables:
    rows_to_keep = []
    is_schema_table = False
    for row in table.rows:
        text = [c.text.lower() for c in row.cells]
        if any("schema" in t for t in text):
            is_schema_table = True
        
        # Check if row is one of the score columns
        if len(text) > 0 and any(col in text[0] for col in ['agent1_score', 'agent2_score', 'unified_score']):
            # Need to remove this row, docx doesn't support easy row deletion, so we delete its XML element
            row._element.getparent().remove(row._element)
        
    if is_schema_table:
        # We assume the schema table has 3 columns (Name, Type, Description) based on output
        # Let's just add to any table that looks like schema
        if len(table.columns) == 3:
            row_cells = table.add_row().cells
            row_cells[0].text = 'raw_complaint_text'
            row_cells[1].text = 'TEXT'
            row_cells[2].text = 'Raw text of customer complaint for NLP'
            
            row_cells = table.add_row().cells
            row_cells[0].text = 'hr_remark_text'
            row_cells[1].text = 'TEXT'
            row_cells[2].text = 'HR remarks for anomaly correlation'

# Update scale text
for p in doc1.paragraphs:
    if "1,000,000" in p.text and "Rows" in p.text:
        p.text = p.text.replace("1,000,000 Rows", "1,000,000 Rows*")
        # Just add it inline or at the end
        if "* 1M for enterprise" not in p.text:
            p.add_run("\n* 1M for enterprise pitch, 50,000 for hackathon live-demo stability").italic = True

for table in doc1.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if "1,000,000" in p.text and ("Rows" in p.text or "rows" in p.text):
                    if "* 1M for enterprise" not in p.text:
                        p.text = p.text.replace("1,000,000 Rows", "1,000,000 Rows*")
                        p.text = p.text.replace("1,000,000 rows", "1,000,000 rows*")
                        p.add_run("\n* 1M for enterprise pitch, 50,000 for hackathon live-demo stability").italic = True


doc1.save('02_Data_Creation_Parameters_Updated.docx')

# 2. Update VaultMind_Complete_Build_Guide.docx
doc2 = docx.Document('VaultMind_Complete_Build_Guide.docx')

# Phase 1 Data Splitting
for p in doc2.paragraphs:
    if "Phase 1" in p.text and "Data" in p.text:
        p.insert_paragraph_before("Data Splitting: The dataset will be split into two parts: Historical (Oct-Feb) and Live Stream (March).", style='List Bullet')
        break

# Phase 3 ML Agents loading history
for p in doc2.paragraphs:
    if "Phase 3" in p.text and "ML" in p.text:
        p.insert_paragraph_before("Backend Graph Init: As soon as the backend starts, it must load 'historical_graph.pkl' so the GNN has prior relational knowledge before the live stream begins.", style='List Bullet')
        break

# Schema update in doc2 as well
for table in doc2.tables:
    is_schema_table = False
    for row in table.rows:
        text = [c.text.lower() for c in row.cells]
        if any("schema" in t for t in text):
            is_schema_table = True
        
        if len(text) > 0 and any(col in text[0] for col in ['agent1_score', 'agent2_score', 'unified_score']):
            row._element.getparent().remove(row._element)
        
    if is_schema_table:
        if len(table.columns) == 3:
            row_cells = table.add_row().cells
            row_cells[0].text = 'raw_complaint_text'
            row_cells[1].text = 'TEXT'
            row_cells[2].text = 'Raw text of customer complaint for NLP'
            
            row_cells = table.add_row().cells
            row_cells[0].text = 'hr_remark_text'
            row_cells[1].text = 'TEXT'
            row_cells[2].text = 'HR remarks for anomaly correlation'

doc2.save('VaultMind_Complete_Build_Guide_Updated.docx')


# 3. Update 05_Testing_Checklist.docx
doc3 = docx.Document('05_Testing_Checklist.docx')

# Add to Section A2
for p in doc3.paragraphs:
    if "Section A2" in p.text or "ML Model Tests" in p.text:
        # Insert after this paragraph
        p.insert_paragraph_before("[ ]  NLP Verification: Verify Agent 4 successfully extracts entities from raw_complaint_text")
        p.insert_paragraph_before("[ ]  Score Calculation Check: Ensure Unified Score is calculated by Backend in real-time, not read from CSV")
        break
else:
    # If couldn't find, just add to the end
    doc3.add_paragraph("[ ]  NLP Verification: Verify Agent 4 successfully extracts entities from raw_complaint_text")
    doc3.add_paragraph("[ ]  Score Calculation Check: Ensure Unified Score is calculated by Backend in real-time, not read from CSV")

doc3.save('05_Testing_Checklist_Updated.docx')

print("All documents updated successfully.")
