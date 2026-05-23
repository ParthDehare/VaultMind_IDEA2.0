import sys
from docx import Document
from docx.shared import Pt, RGBColor

def add_heading(doc, text, level):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(16)
    else:
        run.font.size = Pt(14)
    return p

def add_paragraph(doc, text, bold=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph()
    run = p.add_run('• ' + text)
    return p

doc = Document('VaultMind_Complete_Build_Guide.docx')

add_heading(doc, 'Architecture Addendum: Static Data to Live Stream (The Magic Trick)', level=1)
add_paragraph(doc, "In a hackathon or testing setting, we don't have a live core banking system pumping out real transactions. To simulate a true live environment and prove our real-time streaming architecture works, we use a 'Magic Trick' approach:")

add_heading(doc, '1. The CSV (Static Data)', level=3)
add_paragraph(doc, "Our data generators produce a static CSV dataset (e.g., 1 Lakh rows). We do not let the backend read this file directly.")

add_heading(doc, '2. The Kafka Producer Script (The Simulator)', level=3)
add_paragraph(doc, "We write a Python script (e.g., stream_simulator.py). This script opens the CSV and loops through it row by row. After reading each row, it executes time.sleep(0.5) (a half-second pause) and then pushes that data to the Kafka topic.")

add_heading(doc, '3. The Kafka Broker (The Live Pipeline)', level=3)
add_paragraph(doc, "As far as Kafka is concerned, it is receiving a brand new transaction every 0.5 seconds. For the pipeline, this IS live data!")

add_heading(doc, '4. The Backend Consumer (FastAPI + AI)', level=3)
add_paragraph(doc, "The FastAPI backend consumes the data from Kafka as it arrives, the AI processes it instantly, and the resulting score is pushed via WebSocket to the React UI. The investigator sees a live, ticking dashboard, proving the end-to-end streaming capability without needing a real bank backend.")

doc.save('VaultMind_Complete_Build_Guide.docx')
print("Successfully appended Magic Trick Architecture to VaultMind_Complete_Build_Guide.docx")
